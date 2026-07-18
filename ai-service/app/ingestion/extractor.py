import io
import logging
from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text content from PDF file bytes page-by-page."""
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        text_list = []
        for idx, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_list.append(page_text)
        return "\n\n".join(text_list)
    except Exception as e:
        logger.error("Failed to extract text from PDF: %s", e)
        raise ValueError(f"Failed to parse PDF document: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text content from DOCX paragraphs and tables."""
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        text_list = []

        # Extract standard text paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_list.append(paragraph.text)

        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                # De-duplicate cell values that repeat due to cell merges
                seen = set()
                deduped_row_text = []
                for cell_val in row_text:
                    if cell_val not in seen:
                        seen.add(cell_val)
                        deduped_row_text.append(cell_val)
                if deduped_row_text:
                    text_list.append(" | ".join(deduped_row_text))

        return "\n\n".join(text_list)
    except Exception as e:
        logger.error("Failed to extract text from DOCX: %s", e)
        raise ValueError(f"Failed to parse DOCX document: {str(e)}")


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Unified document parser wrapper routing to correct extractor by filename extension."""
    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format for {filename}. Only PDF and DOCX files are allowed.")
